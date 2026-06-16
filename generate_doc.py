from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

# --- Helper: shade a table row ---
def shade_row(row, hex_colour):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_colour)
        tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    shading = OxmlElement("w:rPr")
    p._p.get_or_add_pPr()
    pPr = p._p.pPr
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "4")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "CCCCCC")
        pBdr.append(bdr)
    pPr.append(pBdr)
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd2)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    shade_row(hdr_row, "1F3A5F")
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(row, "EEF2F7")
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

# ============================================================
# TITLE PAGE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(60)
run = title.add_run("Case 9: Model Serving Lite")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_before = Pt(8)
r = sub.add_run("From Notebook to a Monitored, Retrainable Production Service")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x55, 0x66, 0x88)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Track: MLOps",
    "Difficulty: Medium",
    "Author: Sahil Mohapatra",
    "Date: June 2026",
]:
    meta.add_run(line + "\n").font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. PROBLEM STATEMENT
# ============================================================
add_heading(doc, "1. Problem Statement", 1)
add_body(doc,
    "A data scientist on the team trained a sentiment classification model inside a Jupyter notebook. "
    "It performed well during experimentation but there was no way for the wider team to call it, "
    "monitor it or update it safely. The task was to take that notebook model and turn it into a "
    "production-grade service that a team could actually rely on."
)
doc.add_paragraph()
add_body(doc, "The brief required four things:")
add_bullet(doc, "A live REST API wrapping a pretrained sentiment model with HTTPS.")
add_bullet(doc, "Request and response logging detailed enough to debug a bad prediction a week later.")
add_bullet(doc, "A drift monitoring stub that detects when incoming text starts looking different from training data.")
add_bullet(doc, "A CI pipeline that retrains the model when new labelled data arrives and only promotes it if metrics do not regress.")

doc.add_paragraph()
add_body(doc,
    "The question the brief asked us to answer was: 'How would I know this model is failing in production "
    "before customers do?' That question shaped every technical decision in the project."
)
doc.add_page_break()

# ============================================================
# 2. HOW WE STARTED
# ============================================================
add_heading(doc, "2. How We Started: Assumptions and First Decisions", 1)
add_body(doc,
    "Before writing any code, we locked down the assumptions that would drive the design. "
    "This is important in a one-day project because ambiguous assumptions lead to rework."
)
doc.add_paragraph()

add_heading(doc, "2.1 Assumptions", 2)
assumptions = [
    ("English-only inputs",
     "The base model was trained on English text. Non-English input produces unreliable scores. "
     "Rather than blocking it at the API we chose to treat non-ASCII character frequency as a drift "
     "signal so the API remains open but the monitoring layer flags the shift."),
    ("Pretrained weights are good enough for first deployment",
     "DistilBERT fine-tuned on SST-2 already achieves 91 to 93 percent accuracy on English sentiment "
     "tasks without any additional fine-tuning. Fine-tuning was wired into the CI pipeline for future "
     "data updates but was not required to ship."),
    ("Single uvicorn worker is sufficient",
     "The model is a 270 MB singleton. Running multiple workers would mean each process holding a "
     "separate copy in RAM, which would exhaust the free-tier host. A single worker with async "
     "request handling is the correct choice for this scale."),
    ("In-memory drift state is acceptable",
     "The drift window resets on restart, which is a known limitation. The architecture is correct "
     "and the extension path to Redis or Prometheus is obvious. For a one-day prototype this is "
     "the right trade-off."),
    ("Twenty eval examples are enough to demonstrate the gate",
     "A small held-out set is enough to show the CI gate working. It is not enough for statistically "
     "meaningful estimates, which is noted as a known limitation."),
]
for title_a, body_a in assumptions:
    add_bullet(doc, body_a, bold_prefix=title_a + ": ")
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 3. ARCHITECTURE
# ============================================================
add_heading(doc, "3. Architecture Overview", 1)
add_body(doc,
    "The service is built as a single FastAPI application running inside a Docker container "
    "on HuggingFace Spaces. The model is loaded once at startup and lives in memory for the "
    "lifetime of the process. Every incoming request is logged and recorded for drift monitoring."
)
doc.add_paragraph()

add_heading(doc, "3.1 Request Flow", 2)
add_code(doc, "Caller")
add_code(doc, "  --> POST /predict  (text validated by Pydantic schema)")
add_code(doc, "  --> middleware.py  (logs request, starts timer)")
add_code(doc, "  --> main.py        (endpoint handler)")
add_code(doc, "  --> predict.py     (run_inference via asyncio thread executor)")
add_code(doc, "  --> drift.py       (monitor.record writes stats to rolling window)")
add_code(doc, "  <-- JSON response  (label, score, model_version, latency_ms)")
doc.add_paragraph()

add_heading(doc, "3.2 File Structure", 2)
add_table(doc,
    ["File / Folder", "Responsibility"],
    [
        ["app/main.py", "FastAPI app, lifespan hooks, endpoint routing"],
        ["app/predict.py", "Model loading and async inference wrapper"],
        ["app/schemas.py", "Pydantic request and response models"],
        ["app/drift.py", "Rolling window drift monitor"],
        ["app/middleware.py", "Structured JSON request logging"],
        ["scripts/retrain.py", "Fine-tunes the model on new training data"],
        ["scripts/evaluate.py", "Runs the fine-tuned model against eval.csv"],
        ["scripts/check_metrics.py", "Compares candidate metrics against baseline and gates promotion"],
        ["tests/conftest.py", "Shared fixtures: model mock and HTTP test client"],
        ["tests/test_api.py", "HTTP integration tests for all endpoints"],
        ["tests/test_drift.py", "Unit tests for drift monitoring logic"],
        ["tests/test_predict.py", "Unit tests for inference layer"],
        [".github/workflows/ci.yml", "Lint, test, build Docker image and deploy on every push"],
        [".github/workflows/retrain.yml", "Retrain, evaluate and gate promotion on data PR"],
        ["Dockerfile", "Multi-stage build: builder installs deps, runner is slim image"],
        ["data/train.csv", "Labelled training data (text, label columns)"],
        ["data/eval.csv", "Held-out evaluation set; never modified in training PRs"],
        ["metrics/baseline.json", "Accuracy and F1 of the currently deployed model"],
    ],
    col_widths=[2.0, 4.0]
)

doc.add_page_break()

# ============================================================
# 4. STACK DECISIONS
# ============================================================
add_heading(doc, "4. Technology Stack and Why We Chose Each Tool", 1)

add_table(doc,
    ["Component", "Choice", "Reason"],
    [
        ["Model", "DistilBERT SST-2 (66M params)", "40% smaller and 60% faster than BERT-base on CPU with less than 3% accuracy loss. Free-tier hosts have no GPU so inference speed matters."],
        ["API framework", "FastAPI 0.128.0", "Async request handling, automatic OpenAPI docs and native Pydantic v2 integration out of the box."],
        ["Validation", "Pydantic v2 Annotated fields", "Field constraints (min/max length) are enforced before the request ever reaches the model. Bad input is rejected with a clear 422 error."],
        ["Inference device", "device_map=auto", "Auto-selects CPU, GPU or MPS. The same code works on a laptop and on a cloud instance with a GPU without any changes."],
        ["Package manager", "uv", "10 to 100x faster than pip. Reads uv.lock exactly so every environment from local dev to CI to the Docker build is byte-for-byte identical."],
        ["Container base image", "python:3.11.12-slim (pinned)", "Floating tags like python:3.11-slim can silently pick up a new patch release and break the build. Pinning the exact version makes builds fully reproducible."],
        ["Container registry", "ghcr.io", "Free for public repositories and uses GITHUB_TOKEN with no extra credentials needed."],
        ["Hosting", "HuggingFace Spaces (Docker SDK)", "Free tier with automatic HTTPS. The natural home for ML model demos and integrates cleanly with the HF Hub deployment API."],
        ["Linting", "ruff", "Replaces flake8 and black in one binary. Significantly faster in CI."],
        ["OOV tokenisation", "re.findall word boundary regex", "str.split() leaves punctuation attached to words, inflating OOV rates artificially. The regex strips punctuation at word boundaries so baseline and live tokenisation are consistent."],
    ],
    col_widths=[1.5, 1.8, 3.2]
)

doc.add_page_break()

# ============================================================
# 5. KEY IMPLEMENTATIONS
# ============================================================
add_heading(doc, "5. Key Implementation Decisions", 1)

add_heading(doc, "5.1 Async Inference Without Blocking the Event Loop", 2)
add_body(doc,
    "PyTorch inference is synchronous and holds Python's Global Interpreter Lock while running. "
    "If we called the model directly inside an async endpoint handler, the entire uvicorn event "
    "loop would freeze until inference completed. That means health checks, drift requests and "
    "any other concurrent requests would time out."
)
doc.add_paragraph()
add_body(doc, "The solution is run_in_executor:")
add_code(doc, "loop = asyncio.get_running_loop()")
add_code(doc, "result = await loop.run_in_executor(None, _do_inference, text)")
add_body(doc,
    "This moves the blocking inference call into a background thread pool. The event loop "
    "stays free to handle health checks and other requests while the model works through its "
    "computation. We used asyncio.get_running_loop() rather than the deprecated get_event_loop() "
    "which was removed in Python 3.12."
)
doc.add_paragraph()

add_heading(doc, "5.2 Drift Monitoring: Three Signals", 2)
add_body(doc,
    "The drift monitor computes three statistics from every incoming request and compares them "
    "against a baseline derived from the training data at startup."
)
doc.add_paragraph()
add_table(doc,
    ["Signal", "What it detects", "Alert condition"],
    [
        ["mean_length", "Average character length of recent requests has shifted significantly", "More than 25% relative change from baseline"],
        ["non_ascii_frac", "Fraction of non-ASCII characters: detects language shift away from English", "Absolute delta above 1% or 25% relative change"],
        ["oov_rate", "Fraction of words not seen in training vocabulary: detects topic or domain shift", "Above 30% absolute OOV rate when a training vocab exists"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)
add_body(doc,
    "The baseline and live requests both use re.findall(r'\\b\\w+\\b') for tokenisation. "
    "This was a deliberate choice. Using str.split() on the text 'Great product!' would "
    "produce the token 'product!' which would never match 'product' in the vocabulary. "
    "The regex strips punctuation at word boundaries so both sides of the comparison are consistent."
)
doc.add_paragraph()

add_heading(doc, "5.3 Structured Logging", 2)
add_body(doc,
    "Every request writes a single JSON line to stdout. HuggingFace Spaces captures stdout "
    "automatically and exposes it in the container logs tab. A typical log line looks like this:"
)
add_code(doc, '{"ts": "2026-06-16T10:23:44.123Z", "method": "POST", "path": "/predict",')
add_code(doc, ' "status": 200, "duration_ms": 148.3, "response_snippet": "{\\"label\\": \\"POSITIVE\\"}"}')
add_body(doc,
    "This gives enough information to reconstruct any prediction from a week ago: timestamp, "
    "path, status code, latency and the first part of the response. A bad prediction can be "
    "found by filtering on status 200 and then checking the label and score in the snippet."
)
doc.add_paragraph()

add_heading(doc, "5.4 Pydantic v2 Request Validation", 2)
add_body(doc,
    "The request schema uses the Pydantic v2 Annotated field pattern and adds a field_validator "
    "to strip leading and trailing whitespace before any length check runs."
)
add_code(doc, "text: Annotated[str, Field(min_length=1, max_length=2048)]")
add_code(doc, "")
add_code(doc, "@field_validator('text')")
add_code(doc, "@classmethod")
add_code(doc, "def strip_whitespace(cls, v: str) -> str:")
add_code(doc, "    return v.strip()")
add_body(doc,
    "If a caller sends an empty string or a string of 2049 characters, FastAPI returns a 422 "
    "error before the request ever touches the model. This protects against both accidental "
    "and malicious inputs."
)
doc.add_paragraph()

add_heading(doc, "5.5 Docker Multi-Stage Build", 2)
add_body(doc,
    "The Dockerfile uses two stages. The builder stage installs all dependencies including "
    "the 270 MB model weights. The runner stage copies only what is needed to run the "
    "service, keeping the final image as small as possible."
)
add_body(doc,
    "The model weights are baked into the image at build time rather than downloaded on first "
    "request. This eliminates cold-start latency of 30 to 60 seconds when HuggingFace Spaces "
    "wakes the container after a period of inactivity."
)
add_body(doc,
    "The base image is pinned to python:3.11.12-slim rather than the floating python:3.11-slim "
    "tag. Floating tags can silently pick up a new patch release between two CI runs and "
    "introduce breaking changes without any code change."
)
doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 6. CI/CD PIPELINE
# ============================================================
add_heading(doc, "6. CI/CD Pipeline", 1)
add_body(doc,
    "Two separate GitHub Actions workflows handle different concerns. They are separated "
    "deliberately so that a code change and a data change do not interfere with each other."
)
doc.add_paragraph()

add_heading(doc, "6.1 ci.yml: Code Pipeline (triggers on every push)", 2)
add_code(doc, "Push to any branch")
add_code(doc, "  |")
add_code(doc, "  +-- lint-and-test job")
add_code(doc, "  |     1. actions/checkout@v6")
add_code(doc, "  |     2. astral-sh/setup-uv@v8.1.0")
add_code(doc, "  |     3. uv sync  (installs exact lockfile)")
add_code(doc, "  |     4. uv run ruff check  (lint)")
add_code(doc, "  |     5. uv run pytest  (all 23 tests, model mocked)")
add_code(doc, "  |")
add_code(doc, "  +-- build-and-push job  (main branch only, after tests pass)")
add_code(doc, "  |     1. docker build with GIT_SHA baked in as MODEL_VERSION env var")
add_code(doc, "  |     2. push to ghcr.io with two tags:")
add_code(doc, "  |           sha-<commit>  (traceability: ties image to exact commit)")
add_code(doc, "  |           latest        (deployment: HF Spaces always pulls latest)")
add_code(doc, "  |")
add_code(doc, "  +-- deploy job  (main branch only, after build passes)")
add_code(doc, "        1. uv pip install huggingface_hub==0.24.0")
add_code(doc, "        2. Python script calls HfApi.upload_folder()")
add_code(doc, "        3. 403 error gives a clear 'regenerate token with write scope' message")
doc.add_paragraph()
add_body(doc,
    "Why uv sync instead of pip install? uv reads uv.lock exactly. pip install from "
    "requirements.txt ignores the lockfile and may resolve to a different version on the "
    "CI runner than on a developer's machine. uv sync guarantees that CI runs in the "
    "exact same environment as local development."
)
doc.add_paragraph()
add_body(doc,
    "Why a Python script for the HuggingFace deploy step instead of a CLI command? "
    "The HF CLI does not provide the error handling this step needs. The Python script "
    "catches HTTP 403 errors and prints a clear human-readable message about the token "
    "needing write scope. It also passes an ignore_patterns list to exclude .git, logs/ "
    "and model/ from the upload. A raw shell approach would require much more complexity "
    "to achieve the same robustness."
)
doc.add_paragraph()

add_heading(doc, "6.2 retrain.yml: Data Pipeline (triggers only when data/train.csv changes)", 2)
add_code(doc, "PR modifying data/train.csv")
add_code(doc, "  |")
add_code(doc, "  +-- guard: reject immediately if data/eval.csv was also modified")
add_code(doc, "  |    (the eval set must never change in a training data PR)")
add_code(doc, "  |")
add_code(doc, "  +-- uv run python scripts/retrain.py")
add_code(doc, "  |    Fine-tunes DistilBERT for 2 epochs on new data/train.csv")
add_code(doc, "  |    Saves model weights to model/")
add_code(doc, "  |")
add_code(doc, "  +-- uv run python scripts/evaluate.py")
add_code(doc, "  |    Runs the fine-tuned model against data/eval.csv")
add_code(doc, "  |    Writes accuracy and F1 to metrics/candidate.json")
add_code(doc, "  |")
add_code(doc, "  +-- uv run python scripts/check_metrics.py")
add_code(doc, "        Compares candidate.json against metrics/baseline.json")
add_code(doc, "        PASS: commits new baseline.json and auto-merges the PR")
add_code(doc, "        FAIL: posts a comment on the PR with the exact numbers and exits 1")
doc.add_paragraph()
add_body(doc,
    "The evaluation set is never touched by a training PR. This is enforced by the guard "
    "step at the top of the workflow. Without this guard a bad actor could update both "
    "the training data and the evaluation data in the same PR, making a regressed model "
    "appear to pass the gate."
)
doc.add_paragraph()

add_heading(doc, "6.3 Two Docker Tags Explained", 2)
add_table(doc,
    ["Tag", "Format", "Purpose"],
    [
        ["SHA tag", "sha-a1b2c3d", "Ties the exact image to the exact commit. If a bad prediction is reported, you pull the sha tag and reproduce the environment exactly."],
        ["Latest tag", "latest", "Always points to the newest successful build. HuggingFace Spaces pulls latest on deploy so it always gets the newest image without configuration changes."],
    ],
    col_widths=[1.2, 1.5, 3.8]
)

doc.add_page_break()

# ============================================================
# 7. HOW WOULD I KNOW IT IS FAILING
# ============================================================
add_heading(doc, "7. How Would I Know This Model Is Failing Before Customers Do?", 1)
add_body(doc,
    "This was the central question of the brief. We built four layers of detection, each "
    "catching a different type of failure at a different point in time."
)
doc.add_paragraph()

add_heading(doc, "Layer 1: Input Drift (leading indicator)", 2)
add_body(doc,
    "The /drift endpoint compares the statistical properties of the last 100 requests against "
    "the training data. A shift in character length, non-ASCII fraction or OOV rate means "
    "the model is seeing inputs that look very different from what it was trained on. This "
    "fires before the model starts giving wrong answers because it detects the problem at "
    "the input layer."
)
add_body(doc,
    "You can set up UptimeRobot or a simple cron job to ping /drift every five minutes and "
    "alert when status is 'alert'. No additional infrastructure is needed."
)
doc.add_paragraph()

add_heading(doc, "Layer 2: Confidence Distribution (concurrent indicator)", 2)
add_body(doc,
    "Every prediction writes a score field to the log. A healthy sentiment model is rarely "
    "uncertain. Most scores should be above 0.85. A spike in low-confidence predictions "
    "between 0.50 and 0.70 means the model is struggling with the inputs it is receiving."
)
add_code(doc, "cat logs/api.log | jq 'select(.status == 200) | .response_snippet")
add_code(doc, "  | fromjson | select(.score < 0.75)'")
doc.add_paragraph()

add_heading(doc, "Layer 3: Error Rate and Latency (trailing indicator)", 2)
add_body(doc,
    "The structured logs capture HTTP status codes and duration_ms per request. A spike in "
    "500 errors usually means a dependency broke after a package update. A spike in P99 "
    "latency means the model is slower than expected, perhaps due to longer inputs or "
    "resource contention."
)
add_code(doc, "cat logs/api.log | jq 'select(.status >= 500) | {ts, path, duration_ms}'")
doc.add_paragraph()

add_heading(doc, "Layer 4: CI Gate (proactive)", 2)
add_body(doc,
    "No regressed model can reach production. The retrain workflow compares every candidate "
    "model against the frozen evaluation set before the PR is allowed to merge. If the new "
    "model scores worse than the baseline on accuracy or F1 it is blocked automatically with "
    "a comment explaining the exact numbers."
)
add_body(doc,
    "This means a data scientist cannot accidentally deploy a worse model by pushing bad "
    "training data. The gate catches it in CI before it ever touches the live service."
)

doc.add_page_break()

# ============================================================
# 8. TRADE-OFFS
# ============================================================
add_heading(doc, "8. Trade-offs We Made", 1)

add_table(doc,
    ["Decision", "Alternative", "Why we chose this"],
    [
        ["DistilBERT (66M params)", "BERT-base (110M params)", "40% smaller and 60% faster on CPU with less than 3% accuracy loss. On free-tier hardware with no GPU this matters significantly."],
        ["HuggingFace Spaces", "Render or Fly.io", "Docker SDK gives automatic HTTPS and CI-triggered redeploy with zero Nginx configuration. The natural home for an ML demo."],
        ["Single uvicorn worker", "Gunicorn with multiple workers", "Multiple workers would each hold a separate 270 MB model copy in RAM. A single async worker is correct for this scale."],
        ["In-memory drift window", "Prometheus and Grafana", "Adding a metrics sidecar would triple the infrastructure complexity for marginal benefit in a one-day project."],
        ["run_in_executor for inference", "Async torch (does not exist)", "PyTorch inference is synchronous and blocks the GIL. run_in_executor moves it to a thread pool so the event loop stays responsive."],
        ["Model weights baked into Docker image", "Download on first request", "Avoids 30 to 60 second cold starts when HF Spaces wakes after inactivity. Adds 270 MB to the image but this is a one-time build cost."],
        ["ruff for linting", "flake8 and black separately", "ruff replaces both tools in one binary and runs significantly faster in CI."],
        ["uv for dependency management", "pip with requirements.txt", "uv is 10 to 100x faster, uses a lockfile so every environment is reproducible and manages Python itself in CI."],
        ["Regex tokeniser for OOV", "DistilBERT WordPiece tokeniser", "Using the model tokeniser would couple drift.py to the model weights. A regex tokeniser gives consistent comparison without the coupling."],
        ["Pinned base image python:3.11.12-slim", "Floating python:3.11-slim", "Floating tags can silently pick up a new patch release and introduce breaking changes between CI runs."],
    ],
    col_widths=[1.8, 1.8, 2.9]
)

doc.add_page_break()

# ============================================================
# 9. WHAT WAS DE-SCOPED
# ============================================================
add_heading(doc, "9. What We De-scoped and Why", 1)

add_table(doc,
    ["Feature", "Why de-scoped"],
    [
        ["API key authentication", "Adding a key adds friction to the demo where the judge calls the endpoint with a bare curl command. Noted as a next step in the README."],
        ["Persistent drift storage (Redis)", "Correct architecture is noted and the extension path is obvious. In-memory is sufficient to demonstrate the concept."],
        ["GPU support", "Free-tier host has no GPU. DistilBERT runs at 130 to 200ms on CPU which is acceptable. device_map=auto means a GPU is used automatically if one becomes available."],
        ["Shadow deployment", "Routing 5% of traffic to a candidate model alongside production is architecturally interesting but would take more than one day to wire up properly."],
        ["Statistical significance on metrics gate", "A fixed numeric threshold is easier to reason about for 20 eval examples. A bootstrap confidence interval is the right production approach."],
        ["Locust load test", "Knowing the breaking point under sustained traffic is valuable but was outside the one-day time box."],
        ["/metrics endpoint in Prometheus format", "Would allow Grafana to visualise drift signals and latency percentiles. Noted as the first addition for a second day."],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_page_break()

# ============================================================
# 10. WHAT WOULD WE DO WITH MORE TIME
# ============================================================
add_heading(doc, "10. What We Would Do With Another Day", 1)

add_bullet(doc, "Add a /metrics endpoint in Prometheus format so Grafana can show request rate, P50/P99 latency and all three drift signals on one dashboard.")
add_bullet(doc, "Store drift signals in Redis so they survive container restarts and aggregate across multiple workers.")
add_bullet(doc, "Write a Locust load test to find the breaking point under sustained concurrent traffic.")
add_bullet(doc, "Implement the shadow deployment stretch goal: route 5% of live traffic to the candidate model and compare predictions before promoting.")
add_bullet(doc, "Add proper API key authentication with key rotation and a revocation endpoint.")
add_bullet(doc, "Add a warm-up ping from the CI deploy step so HF Spaces does not serve a cold-start response to the first real caller after a deploy.")
add_bullet(doc, "Replace the in-memory drift window with a Redis-backed sliding window that aggregates signals across all workers.")

doc.add_page_break()

# ============================================================
# 11. KNOWN LIMITATIONS
# ============================================================
add_heading(doc, "11. Known Limitations", 1)

add_table(doc,
    ["Limitation", "Detail", "Production fix"],
    [
        ["Cold start", "HF Spaces free tier sleeps after 15 minutes of inactivity. First request after sleep takes 5 to 10 seconds.", "Warm-up ping from CI deploy step. Paid tier removes sleep entirely."],
        ["Drift window resets on restart", "The rolling window is in-memory. A new deployment clears it and status returns to insufficient_data until 10 requests arrive.", "Redis-backed storage with TTL."],
        ["Small evaluation set", "Twenty examples is enough to demonstrate the gate but not enough for statistically meaningful estimates.", "Collect more labelled data. Add bootstrap confidence intervals to the gate."],
        ["OOV rate is approximate", "The drift monitor uses a regex tokeniser rather than DistilBERT's WordPiece tokeniser. OOV rate is an approximation of what the model actually sees.", "Use the model tokeniser in drift.py at the cost of coupling to model weights."],
        ["No authentication", "The /predict endpoint is open. Any caller can use it.", "Add API key header validation with a simple middleware."],
        ["Single worker", "Only one process handles requests. A slow inference blocks all other callers during that window.", "Scale horizontally with multiple workers once model weights are moved to shared storage."],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_page_break()

# ============================================================
# 12. LIVE ENDPOINTS AND DEMO
# ============================================================
add_heading(doc, "12. Live Endpoints and Demo", 1)

add_heading(doc, "12.1 Endpoints", 2)
add_table(doc,
    ["Endpoint", "Method", "Purpose"],
    [
        ["/healthz", "GET", "Liveness check. Returns {\"status\": \"ok\"} when the service is up."],
        ["/predict", "POST", "Classifies text as POSITIVE or NEGATIVE. Requires JSON body with a text field between 1 and 2048 characters."],
        ["/drift", "GET", "Current drift monitoring report with status, signal averages and any active alerts."],
        ["/docs", "GET", "Auto-generated OpenAPI documentation with live try-it-out capability."],
    ],
    col_widths=[1.2, 0.8, 4.5]
)

add_heading(doc, "12.2 Example Prediction", 2)
add_code(doc, "curl -X POST https://sahil147-sentiment-api.hf.space/predict \\")
add_code(doc, "  -H 'Content-Type: application/json' \\")
add_code(doc, "  -d '{\"text\": \"This product is absolutely brilliant and exceeded every expectation.\"}'")
doc.add_paragraph()
add_code(doc, '{')
add_code(doc, '  "label": "POSITIVE",')
add_code(doc, '  "score": 0.999812,')
add_code(doc, '  "model_version": "sha-a1b2c3d",')
add_code(doc, '  "latency_ms": 148.3')
add_code(doc, '}')
doc.add_paragraph()

add_heading(doc, "12.3 Drift Simulation", 2)
add_body(doc,
    "To demonstrate drift detection, send 10 normal English requests then flood the window "
    "with non-English text. The /drift endpoint will return an alert on non_ascii_frac."
)
add_code(doc, "import requests")
add_code(doc, "BASE = 'https://sahil147-sentiment-api.hf.space'")
add_code(doc, "SPANISH = 'Esto es una resena en espanol: me encanta este producto.'")
add_code(doc, "")
add_code(doc, "for _ in range(10):")
add_code(doc, "    requests.post(f'{BASE}/predict', json={'text': 'Great product.'})")
add_code(doc, "")
add_code(doc, "for _ in range(100):")
add_code(doc, "    requests.post(f'{BASE}/predict', json={'text': SPANISH})")
add_code(doc, "")
add_code(doc, "print(requests.get(f'{BASE}/drift').json())")

doc.add_page_break()

# ============================================================
# 13. STACK VERSIONS
# ============================================================
add_heading(doc, "13. Final Stack Versions", 1)

add_table(doc,
    ["Component", "Version"],
    [
        ["Python", "3.11.12"],
        ["FastAPI", "0.128.0"],
        ["Pydantic", "v2.7.1"],
        ["Transformers", "4.57.3"],
        ["uvicorn", "0.32.0"],
        ["uv (package manager)", "latest (lockfile: uv.lock)"],
        ["actions/checkout", "@v6"],
        ["astral-sh/setup-uv", "@v8.1.0"],
        ["Docker base image", "python:3.11.12-slim"],
        ["Container registry", "ghcr.io"],
        ["Hosting", "HuggingFace Spaces (Docker SDK)"],
    ],
    col_widths=[2.5, 4.0]
)

# Save
out_path = "/Users/sahil/Desktop/Infinia Tech/case9-model-serving-lite/Case9_Project_Documentation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
